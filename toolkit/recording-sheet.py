# Build the native-speaker recording sheet (PDF) from content.js.
#
#   python toolkit/recording-sheet.py
#
# Output:
#   toolkit/recording-sheet.pdf   — hand this to the speaker
#   toolkit/recording-sheet.docx  — same content, editable (Megan asked for it 2026-07-18)
#   toolkit/recording-list.json   — machine-readable order (n -> itemId), for tagging later
#
# BOTH documents are built from the same rows() in one run, so the Word copy cannot drift
# from the PDF. But note: once Megan EDITS the .docx, that stops being true — if she adds,
# removes or reorders words, the numbering no longer matches recording-list.json or the
# appendix, and tagging will mis-map. Re-derive the mapping from her edited copy in that case.
#
# The list is the SAME data missing-audio.py reports, merged into one continuous reading
# order: the 68 silent cards plus the 64 flagged for re-recording (132 total). The 8 reuse
# cards are excluded on purpose — they play another card's clip and are fixed for free.
#
# Reading order == recording order == segment order after slicing. Keep them in step: if
# the speaker skips an item, the numbering after it no longer lines up with the segments,
# which is why the sheet asks her not to skip.
#
# Recording guidance in the sheet is derived from slice-lessons.py's real thresholds
# (NOISE_DB -30dB, MIN_SILENCE 0.35s, MIN_SEG 0.4s) — do not soften it without changing
# those. Megan's own recordings measured -29..-31 dB BETWEEN words, i.e. right at the
# detection threshold, which is why Food.mp3 produced merged segments needing manual
# splitting. A quiet room and a generous pause are what make the slicer reliable.

import importlib.util
import json
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
TOOLKIT = Path(__file__).resolve().parent

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

spec = importlib.util.spec_from_file_location('ma', TOOLKIT / 'missing-audio.py')
ma = importlib.util.module_from_spec(spec)
spec.loader.exec_module(ma)


def collect():
    items = ma.parse_content()
    managed = ma.mapping_managed_ids()
    silent = [r for r in items if r[4]['audio'] is None]
    external = [r for r in items if r[4]['audio'] and r[4]['id'] not in managed]
    own = [r for r in external if r[4]['audio'] == f"items/{r[4]['id']}.mp3"]
    rows = [(r, 'new') for r in silent] + [(r, 'redo') for r in own]

    def order(x):
        (uid, _ut, lid, _lt, it), _kind = x
        return (int(uid[1:]), int(lid.split('l')[1]), int(it['id'].split('-')[1]))

    rows.sort(key=order)
    return [{
        'n': i + 1, 'id': r[4]['id'], 'tsw': r[4]['tsw'], 'eng': r[4]['eng'],
        'unit': r[1], 'lesson_id': r[2], 'lesson': r[3], 'kind': kind,
    } for i, (r, kind) in enumerate(rows)]


def tex_escape(s):
    for a, b in (('\\', r'\textbackslash{}'), ('&', r'\&'), ('%', r'\%'), ('$', r'\$'),
                 ('#', r'\#'), ('_', r'\_'), ('{', r'\{'), ('}', r'\}'),
                 ('~', r'\textasciitilde{}'), ('^', r'\textasciicircum{}')):
        s = s.replace(a, b)
    return s


PREAMBLE = r"""
\documentclass[11pt,a4paper]{article}
\usepackage[margin=1.9cm,top=1.7cm,bottom=1.9cm]{geometry}
\usepackage{fontspec}
\usepackage{longtable}
\usepackage{array}
\usepackage[table]{xcolor}  % [table] = \rowcolor; must load before tcolorbox pulls xcolor in
\usepackage{fancyhdr}
\usepackage{tcolorbox}
\usepackage{enumitem}
\setmainfont{Segoe UI}
\definecolor{ink}{HTML}{1F2937}
\definecolor{accent}{HTML}{7C3AED}
\definecolor{soft}{HTML}{F3F0FF}
\definecolor{rule}{HTML}{D8D2EC}
\definecolor{muted}{HTML}{6B7280}
\color{ink}
\setlength{\parindent}{0pt}
\renewcommand{\arraystretch}{1.35}
\pagestyle{fancy}\fancyhf{}
\renewcommand{\headrulewidth}{0pt}
\fancyfoot[C]{\small\color{muted}Re:Lefela recording sheet \quad·\quad page \thepage}
"""


def build_tex(rows):
    L = [PREAMBLE, r'\begin{document}', '']
    L.append(r'{\LARGE\bfseries\color{accent} Setswana recordings for Re:Lefela}\\[4pt]')
    L.append(r'{\color{muted}A word list to read aloud \quad·\quad %d items \quad·\quad about 10--15 minutes}'
             % len(rows))
    L.append(r'\vspace{10pt}')

    L.append(r"""
\begin{tcolorbox}[colback=soft,colframe=rule,boxrule=0.8pt,arc=3pt,left=9pt,right=9pt,top=7pt,bottom=7pt]
\textbf{Thank you so much for doing this.} Re:Lefela is a little app my study partner and I
built to learn Setswana properly. Every word below is one the app teaches but has no good
native recording for yet --- right now it uses my own voice, which is not what anyone should
be learning pronunciation from.

\medskip
\textbf{You do not need to record these one at a time.} Just make \textbf{one long recording}
and read straight down the list. My software finds the silences between words and cuts them
apart automatically.
\end{tcolorbox}
""")

    L.append(r'\vspace{6pt}')
    L.append(r'{\large\bfseries\color{accent} How to record}')
    L.append(r'\vspace{-2pt}')
    L.append(r"""
\begin{itemize}[leftmargin=15pt,itemsep=3.5pt,topsep=6pt]
  \item \textbf{A quiet room matters more than a good microphone.} A phone voice recorder is
        perfectly fine. What breaks the automatic cutting is background noise --- a TV, music,
        a fan, an open window --- because then the gaps between words are not really silent
        and the software cannot find them.
  \item \textbf{Leave about two seconds between each word.} Say the word, then count
        \emph{one\ldots two} in your head, then say the next one. A generous gap is the single
        most helpful thing you can do; too short is the one mistake that is hard for me to fix.
  \item \textbf{Say only the Setswana.} Please do not read the English out loud --- it doubles
        the number of pieces I have to sort through afterwards.
  \item \textbf{Say each word once,} at a normal, unhurried pace. Short words like \emph{mae}
        and \emph{koko} especially --- please do not clip them short.
  \item \textbf{Go in order and please do not skip any.} The numbering is how I match each
        recording back to the right word in the app. If you would rather not say something,
        say the number out loud instead (``twenty-nine'') so I can see the gap.
  \item \textbf{If you fluff one, just pause and say it again.} No need to start over --- I
        keep the last version of anything said twice.
  \item \textbf{Take breaks whenever you like.} A long pause does no harm at all. If you would
        rather record it over a few sittings, that is completely fine --- just tell me where
        you stopped.
\end{itemize}
""")

    L.append(r'\vspace{2pt}')
    L.append(r'{\large\bfseries\color{accent} If something looks wrong}')
    L.append(r"""
\vspace{2pt}

I built this list from a Peace Corps course and a few published sources, so some of it may be
old-fashioned, spelled the South African way rather than the Botswana way, or simply wrong.
\textbf{Please do not quietly correct it as you read} --- if I do not know a word changed, the
recording and the app will disagree with each other and the app will teach the wrong thing.

\medskip
Instead: \textbf{say it the way the list has it}, then use the \textbf{Notes} column to tell me
what it should be. If an English translation is off, note that too. I would genuinely rather
find out now than have us learning it wrong all semester.
""")

    L.append(r'\clearpage')
    L.append(r'{\large\bfseries\color{accent} The word list}\\[2pt]')
    L.append(r'{\color{muted}\small Read straight down. Setswana only --- the English is just so you can check it.}')
    L.append(r'\vspace{8pt}')

    header = (r'\rowcolor{soft}\textbf{\#} & \textbf{Setswana --- say this} & '
              r'\textbf{English (check me)} & \textbf{Notes} \\')
    # m{} (not p{}) — vertically CENTRES each cell. With p{} the columns top-align on their
    # own first baseline, and because the Setswana column is 11pt while # and English are
    # \small, the three cells of one row visibly sit on different lines.
    L.append(r'\begin{longtable}{@{}>{\color{muted}\small}m{0.9cm}'
             r'>{\raggedright\arraybackslash}m{5.1cm}'
             r'>{\raggedright\arraybackslash\color{muted}}m{4.7cm}'
             r'>{\raggedright\arraybackslash}m{4.4cm}@{}}')
    L.append(r'\hline ' + header + r'\hline')
    L.append(r'\endfirsthead')
    L.append(r'\hline ' + header + r'\hline')
    L.append(r'\endhead')

    last = None
    for r in rows:
        if r['lesson_id'] != last:
            last = r['lesson_id']
            L.append(r'\multicolumn{4}{@{}l@{}}{\vspace{2pt}} \\[-8pt]')
            # Lesson title only. The unit title used to sit alongside it, but Unit 1's is
            # "Go dumedisa!", so four consecutive lessons all read "... Go dumedisa!" and it
            # looked like a bug. The speaker does not need the app's unit structure anyway.
            L.append(r'\multicolumn{4}{@{}l@{}}{\textbf{\color{accent}%s}} \\[2pt]'
                     % tex_escape(r['lesson']))
        L.append(r'%d & %s & %s & \\' % (r['n'], tex_escape(r['tsw']), tex_escape(r['eng'])))
        L.append(r'\arrayrulecolor{rule}\hline')
    L.append(r'\end{longtable}')

    L.append(r'\vspace{4pt}')
    L.append(r'{\color{muted}\small That is everything --- thank you again. '
             r'Any format is fine: WhatsApp voice note, phone recording, whatever is easiest.}')

    # Appendix: Megan-only mapping, deliberately after the speaker-facing content.
    L.append(r'\clearpage')
    L.append(r'{\large\bfseries\color{accent} For Megan --- tagging reference}\\[2pt]')
    L.append(r'{\color{muted}\small Not for the speaker. Recording order $\rightarrow$ card id. '
             r'\textbf{new} = first recording (new file). \textbf{redo} = overwrites an existing '
             r'\texttt{audio/items/<id>.mp3}, so this wave WILL need \texttt{AUDIO\_CACHE} '
             r'(\texttt{relefela-audio-vN}) bumped in \texttt{sw.js}. '
             r'Same data as \texttt{toolkit/recording-list.json}.}')
    L.append(r'\vspace{8pt}')
    L.append(r'{\footnotesize')
    L.append(r'\begin{longtable}{@{}>{\small}p{0.8cm}p{2.3cm}p{4.0cm}>{\small}p{0.8cm}p{2.3cm}p{4.0cm}@{}}')
    half = (len(rows) + 1) // 2
    left, right = rows[:half], rows[half:]
    for i in range(half):
        a = left[i]
        b = right[i] if i < len(right) else None
        cells = [str(a['n']), tex_escape(a['id']) + r'\,{\color{muted}' + a['kind'] + '}', tex_escape(a['tsw'])]
        if b:
            cells += [str(b['n']), tex_escape(b['id']) + r'\,{\color{muted}' + b['kind'] + '}', tex_escape(b['tsw'])]
        else:
            cells += ['', '', '']
        L.append(' & '.join(cells) + r' \\')
    L.append(r'\end{longtable}}')

    L.append(r'\end{document}')
    return '\n'.join(L)


INTRO_1 = ('Thank you so much for doing this. Re:Lefela is a little app my study partner and I '
           'built to learn Setswana properly. Every word below is one the app teaches but has no '
           'good native recording for yet — right now it uses my own voice, which is not what '
           'anyone should be learning pronunciation from.')
INTRO_2 = ('You do not need to record these one at a time. Just make one long recording and read '
           'straight down the list. My software finds the silences between words and cuts them '
           'apart automatically.')

HOW_TO = [
    ('A quiet room matters more than a good microphone.',
     ' A phone voice recorder is perfectly fine. What breaks the automatic cutting is background '
     'noise — a TV, music, a fan, an open window — because then the gaps between words are not '
     'really silent and the software cannot find them.'),
    ('Leave about two seconds between each word.',
     ' Say the word, then count one… two in your head, then say the next one. A generous gap is '
     'the single most helpful thing you can do; too short is the one mistake that is hard for me '
     'to fix.'),
    ('Say only the Setswana.',
     ' Please do not read the English out loud — it doubles the number of pieces I have to sort '
     'through afterwards.'),
    ('Say each word once,',
     ' at a normal, unhurried pace. Short words like mae and koko especially — please do not clip '
     'them short.'),
    ('Go in order and please do not skip any.',
     ' The numbering is how I match each recording back to the right word in the app. If you would '
     'rather not say something, say the number out loud instead ("twenty-nine") so I can see the gap.'),
    ('If you fluff one, just pause and say it again.',
     ' No need to start over — I keep the last version of anything said twice.'),
    ('Take breaks whenever you like.',
     ' A long pause does no harm at all. If you would rather record it over a few sittings, that is '
     'completely fine — just tell me where you stopped.'),
]

WRONG_1 = ('I built this list from a Peace Corps course and a few published sources, so some of it '
           'may be old-fashioned, spelled the South African way rather than the Botswana way, or '
           'simply wrong. Please do not quietly correct it as you read — if I do not know a word '
           'changed, the recording and the app will disagree with each other and the app will teach '
           'the wrong thing.')
WRONG_2 = ('Instead: say it the way the list has it, then use the Notes column to tell me what it '
           'should be. If an English translation is off, note that too. I would genuinely rather '
           'find out now than have us learning it wrong all semester.')


def build_docx(rows, path):
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    ACCENT = RGBColor(0x7C, 0x3A, 0xED)
    MUTED = RGBColor(0x6B, 0x72, 0x80)
    INK = RGBColor(0x1F, 0x29, 0x37)

    def shade(cell, hex_fill):
        el = OxmlElement('w:shd')
        el.set(qn('w:val'), 'clear')          # never 'solid' — renders black
        el.set(qn('w:fill'), hex_fill)
        cell._tc.get_or_add_tcPr().append(el)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    for m in ('left_margin', 'right_margin'):
        setattr(sec, m, Cm(1.9))
    sec.top_margin, sec.bottom_margin = Cm(1.7), Cm(1.9)

    normal = doc.styles['Normal']
    normal.font.name = 'Segoe UI'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    # east-asian mapping too, else Word substitutes a different face for some runs
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Segoe UI')

    def heading(text, size=14):
        # Real 'Heading 1' style, not a bold Normal paragraph — gives Megan the Navigation
        # pane and lets her restyle every heading at once, which is the point of a .docx.
        p = doc.add_paragraph(style='Heading 1')
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        r.bold = True
        r.font.name = 'Segoe UI'
        r.font.size = Pt(size)
        r.font.color.rgb = ACCENT
        return p

    def body(*parts):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        for text, bold in parts:
            r = p.add_run(text)
            r.bold = bold
        return p

    title = doc.add_paragraph(style='Title')
    title.paragraph_format.space_after = Pt(2)
    tr = title.add_run('Setswana recordings for Re:Lefela')
    tr.font.name = 'Segoe UI'
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(12)
    sr = sub.add_run(f'A word list to read aloud  ·  {len(rows)} items  ·  about 10–15 minutes')
    sr.font.color.rgb = MUTED

    body((INTRO_1.split('.')[0] + '.', True), (INTRO_1.split('.', 1)[1], False))
    body(('You do not need to record these one at a time.', True),
         (INTRO_2.split('at a time.', 1)[1], False))

    heading('How to record')
    for lead, rest in HOW_TO:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(5)
        p.add_run(lead).bold = True
        p.add_run(rest)

    heading('If something looks wrong')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.add_run(WRONG_1.split('Please do not')[0])
    p.add_run('Please do not quietly correct it as you read').bold = True
    p.add_run(WRONG_1.split('as you read', 1)[1])
    p = doc.add_paragraph()
    p.add_run('Instead: ')
    p.add_run('say it the way the list has it').bold = True
    p.add_run(', then use the ')
    p.add_run('Notes').bold = True
    p.add_run(' column to tell me what it should be. If an English translation is off, note that '
              'too. I would genuinely rather find out now than have us learning it wrong all semester.')

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading('The word list')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Read straight down. Setswana only — the English is just so you can check it.')
    r.font.color.rgb = MUTED

    widths = (Cm(1.2), Cm(5.4), Cm(5.0), Cm(5.6))
    t = doc.add_table(rows=1, cols=4)
    t.style = 'Table Grid'
    for cell, txt, w in zip(t.rows[0].cells,
                            ('#', 'Setswana — say this', 'English (check me)', 'Notes'), widths):
        cell.width = w
        shade(cell, 'F3F0FF')
        run = cell.paragraphs[0].add_run(txt)
        run.bold = True
    t.rows[0].cells[0].paragraphs[0].runs[0].text = '#'

    last = None
    for r_ in rows:
        if r_['lesson_id'] != last:
            last = r_['lesson_id']
            hdr = t.add_row()
            hdr.cells[0].merge(hdr.cells[3])
            c = hdr.cells[0]
            shade(c, 'FAF8FF')
            run = c.paragraphs[0].add_run(r_['lesson'])
            run.bold = True
            run.font.color.rgb = ACCENT
        row = t.add_row()
        for cell, txt, w in zip(row.cells, (str(r_['n']), r_['tsw'], r_['eng'], ''), widths):
            cell.width = w
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if txt:
                cell.paragraphs[0].add_run(txt)
        row.cells[0].paragraphs[0].runs[0].font.color.rgb = MUTED

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run('That is everything — thank you again. Any format is fine: WhatsApp voice note, '
                  'phone recording, whatever is easiest.')
    r.font.color.rgb = MUTED

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    heading('For Megan — tagging reference')
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Not for the speaker. Recording order → card id. "new" = first recording (new '
                  'file). "redo" = overwrites an existing audio/items/<id>.mp3, so this wave WILL '
                  'need AUDIO_CACHE (relefela-audio-vN) bumped in sw.js. Same data as '
                  'toolkit/recording-list.json. ⚠️ If you edit the word list above — adding, '
                  'removing or reordering anything — this mapping and recording-list.json go stale '
                  'and must be regenerated from the edited list.')
    r.font.color.rgb = MUTED
    r.font.size = Pt(9)

    at = doc.add_table(rows=1, cols=4)
    at.style = 'Table Grid'
    aw = (Cm(1.2), Cm(2.6), Cm(2.2), Cm(11.2))
    for cell, txt, w in zip(at.rows[0].cells, ('#', 'card id', 'kind', 'Setswana'), aw):
        cell.width = w
        shade(cell, 'F3F0FF')
        cell.paragraphs[0].add_run(txt).bold = True
    for r_ in rows:
        row = at.add_row()
        for cell, txt, w in zip(row.cells, (str(r_['n']), r_['id'], r_['kind'], r_['tsw']), aw):
            cell.width = w
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(9)

    doc.save(str(path))


def main():
    rows = collect()
    (TOOLKIT / 'recording-list.json').write_text(
        json.dumps(rows, ensure_ascii=False, indent=1) + '\n', encoding='utf-8', newline='\n')

    tex = TOOLKIT / 'recording-sheet.tex'
    tex.write_text(build_tex(rows), encoding='utf-8', newline='\n')

    for _ in range(2):  # twice: longtable needs a second pass to settle column widths
        p = subprocess.run(['xelatex', '-interaction=nonstopmode', '-halt-on-error',
                            f'-output-directory={TOOLKIT}', str(tex)],
                           capture_output=True, text=True, encoding='utf-8', errors='replace')
        if p.returncode != 0:
            tail = '\n'.join((p.stdout or '').splitlines()[-30:])
            print('xelatex FAILED\n' + tail)
            return 1
    for ext in ('.aux', '.log', '.out'):
        (TOOLKIT / f'recording-sheet{ext}').unlink(missing_ok=True)

    build_docx(rows, TOOLKIT / 'recording-sheet.docx')

    n_new = sum(1 for r in rows if r['kind'] == 'new')
    print(f'{len(rows)} items ({n_new} new, {len(rows) - n_new} re-record)')
    print(f'wrote {TOOLKIT / "recording-sheet.pdf"}')
    print(f'wrote {TOOLKIT / "recording-sheet.docx"}')
    print(f'wrote {TOOLKIT / "recording-list.json"}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
